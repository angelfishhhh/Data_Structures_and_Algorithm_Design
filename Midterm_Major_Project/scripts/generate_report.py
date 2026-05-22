from __future__ import annotations

import csv
import random
import re
import shutil
import subprocess
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
RESULTS = ROOT / "results"
PACKAGE = ROOT / "package"


def run(command: list[str], input_text: str | None = None) -> str:
    completed = subprocess.run(
        command,
        cwd=ROOT,
        input=input_text,
        text=True,
        capture_output=True,
        check=True,
    )
    return completed.stdout


def build_project() -> None:
    shutil.rmtree(ROOT / "build", ignore_errors=True)
    shutil.rmtree(ROOT / "bin", ignore_errors=True)
    run(["cmake", "-S", ".", "-B", "build", "-DCMAKE_BUILD_TYPE=Release"])
    run(["cmake", "--build", "build"])


def parse_selection_output(output: str) -> dict[str, int]:
    result: dict[str, int] = {}
    for name in ["HeapSelect", "RandomizedSelect", "MedianOfMediansSelect"]:
        match = re.search(rf"{name}: (-?\d+) \(time_us: (\d+)\)", output)
        if not match:
            raise ValueError(f"Cannot parse {name} from output:\n{output}")
        result[f"{name}_value"] = int(match.group(1))
        result[f"{name}_time_us"] = int(match.group(2))
    return result


def parse_partition_output(output: str) -> dict[str, int]:
    parsed: dict[str, int] = {}
    for name in ["BruteForce", "DynamicProgramming"]:
        diff = re.search(rf"{name} best difference: (\d+)", output)
        elapsed = re.search(rf"{name} time_us: (\d+)", output)
        if diff:
            parsed[f"{name}_difference"] = int(diff.group(1))
        if elapsed:
            parsed[f"{name}_time_us"] = int(elapsed.group(1))
    return parsed


def make_values(n: int, seed: int) -> list[int]:
    rng = random.Random(seed)
    return [rng.randint(1, 10_000_000) for _ in range(n)]


def selection_input(values: list[int], k: int) -> str:
    return f"{len(values)} {k}\n" + " ".join(map(str, values)) + "\n"


def partition_input(values: list[int]) -> str:
    return " ".join(map(str, values)) + "\n"


def write_csv(path: Path, fieldnames: list[str], rows: list[dict[str, int | str]]) -> None:
    with path.open("w", newline="", encoding="utf-8") as file:
        writer = csv.DictWriter(file, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def draw_line_chart(path: Path, title: str, x_labels: list[str], series: list[tuple[str, list[int]]]) -> None:
    width, height = 1000, 620
    margin_left, margin_right = 90, 40
    margin_top, margin_bottom = 70, 110
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font = ImageFont.load_default()

    plot_left = margin_left
    plot_right = width - margin_right
    plot_top = margin_top
    plot_bottom = height - margin_bottom
    all_values = [value for _, values in series for value in values]
    max_value = max(max(all_values), 1)
    colors = ["#1f77b4", "#d62728", "#2ca02c", "#9467bd"]

    draw.text((width // 2 - 120, 22), title, fill="black", font=font)
    draw.line((plot_left, plot_top, plot_left, plot_bottom), fill="black", width=2)
    draw.line((plot_left, plot_bottom, plot_right, plot_bottom), fill="black", width=2)

    for tick in range(6):
        y = plot_bottom - tick * (plot_bottom - plot_top) / 5
        value = int(max_value * tick / 5)
        draw.line((plot_left - 5, y, plot_left, y), fill="black")
        draw.text((10, y - 6), str(value), fill="black", font=font)
        draw.line((plot_left, y, plot_right, y), fill="#e6e6e6")

    count = len(x_labels)
    x_positions = [
        plot_left + i * (plot_right - plot_left) / max(count - 1, 1)
        for i in range(count)
    ]
    for x, label in zip(x_positions, x_labels):
        draw.line((x, plot_bottom, x, plot_bottom + 5), fill="black")
        draw.text((x - 18, plot_bottom + 12), label, fill="black", font=font)

    for idx, (name, values) in enumerate(series):
        color = colors[idx % len(colors)]
        points = []
        for x, value in zip(x_positions, values):
            y = plot_bottom - value / max_value * (plot_bottom - plot_top)
            points.append((x, y))
        if len(points) > 1:
            draw.line(points, fill=color, width=3)
        for x, y in points:
            draw.ellipse((x - 4, y - 4, x + 4, y + 4), fill=color)
        legend_y = plot_bottom + 42 + idx * 18
        draw.rectangle((plot_left + 10, legend_y, plot_left + 24, legend_y + 10), fill=color)
        draw.text((plot_left + 32, legend_y - 2), name, fill="black", font=font)

    image.save(path)


def generate_experiments() -> tuple[list[dict[str, int | str]], list[dict[str, int | str]], list[dict[str, int | str]]]:
    RESULTS.mkdir(exist_ok=True)

    by_k_rows: list[dict[str, int | str]] = []
    fixed_n = 20_000
    values = make_values(fixed_n, 101)
    for k in [1, 50, 500, 5_000, 10_000, 20_000]:
        output = run([str(ROOT / "bin" / "selection")], selection_input(values, k))
        parsed = parse_selection_output(output)
        by_k_rows.append(
            {
                "n": fixed_n,
                "k": k,
                "heap_us": parsed["HeapSelect_time_us"],
                "randomized_us": parsed["RandomizedSelect_time_us"],
                "median_of_medians_us": parsed["MedianOfMediansSelect_time_us"],
            }
        )

    by_n_rows: list[dict[str, int | str]] = []
    for n in [5_000, 10_000, 20_000, 40_000, 80_000]:
        k = n // 2
        values = make_values(n, 202 + n)
        output = run([str(ROOT / "bin" / "selection")], selection_input(values, k))
        parsed = parse_selection_output(output)
        by_n_rows.append(
            {
                "n": n,
                "k": k,
                "heap_us": parsed["HeapSelect_time_us"],
                "randomized_us": parsed["RandomizedSelect_time_us"],
                "median_of_medians_us": parsed["MedianOfMediansSelect_time_us"],
            }
        )

    partition_rows: list[dict[str, int | str]] = []
    rng = random.Random(303)
    for n in [8, 12, 16, 20, 22]:
        values = [rng.randint(1, 30) for _ in range(n)]
        output = run([str(ROOT / "bin" / "partition")], partition_input(values))
        parsed = parse_partition_output(output)
        partition_rows.append(
            {
                "n": n,
                "total_value": sum(values),
                "brute_force_us": parsed["BruteForce_time_us"],
                "dynamic_programming_us": parsed["DynamicProgramming_time_us"],
                "best_difference": parsed["DynamicProgramming_difference"],
            }
        )

    write_csv(
        RESULTS / "selection_by_k.csv",
        ["n", "k", "heap_us", "randomized_us", "median_of_medians_us"],
        by_k_rows,
    )
    write_csv(
        RESULTS / "selection_by_n.csv",
        ["n", "k", "heap_us", "randomized_us", "median_of_medians_us"],
        by_n_rows,
    )
    write_csv(
        RESULTS / "partition_experiments.csv",
        ["n", "total_value", "brute_force_us", "dynamic_programming_us", "best_difference"],
        partition_rows,
    )

    draw_line_chart(
        RESULTS / "selection_by_k.png",
        "Selection time vs k (n=20000)",
        [str(row["k"]) for row in by_k_rows],
        [
            ("Heap", [int(row["heap_us"]) for row in by_k_rows]),
            ("Randomized", [int(row["randomized_us"]) for row in by_k_rows]),
            ("MedianOfMedians", [int(row["median_of_medians_us"]) for row in by_k_rows]),
        ],
    )
    draw_line_chart(
        RESULTS / "selection_by_n.png",
        "Selection time vs n (k=n/2)",
        [str(row["n"]) for row in by_n_rows],
        [
            ("Heap", [int(row["heap_us"]) for row in by_n_rows]),
            ("Randomized", [int(row["randomized_us"]) for row in by_n_rows]),
            ("MedianOfMedians", [int(row["median_of_medians_us"]) for row in by_n_rows]),
        ],
    )
    draw_line_chart(
        RESULTS / "partition_experiments.png",
        "Partition time comparison",
        [str(row["n"]) for row in partition_rows],
        [
            ("BruteForce", [int(row["brute_force_us"]) for row in partition_rows]),
            ("DynamicProgramming", [int(row["dynamic_programming_us"]) for row in partition_rows]),
        ],
    )

    return by_k_rows, by_n_rows, partition_rows


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def add_code_block(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run_obj = paragraph.add_run(text.strip())
    run_obj.font.name = "Courier New"
    run_obj.font.size = Pt(8)


def set_document_fonts(document: Document) -> None:
    styles = document.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), fill)
    tc_pr.append(shade)


def add_report_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
        set_cell_shading(table.rows[0].cells[index], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def generate_report(
    by_k_rows: list[dict[str, int | str]],
    by_n_rows: list[dict[str, int | str]],
    partition_rows: list[dict[str, int | str]],
) -> None:
    sample_selection = run([str(ROOT / "bin" / "selection")], (ROOT / "tests" / "selection_sample.txt").read_text())
    sample_partition = run([str(ROOT / "bin" / "partition")], (ROOT / "tests" / "partition_sample.txt").read_text())
    sample_mst = run([str(ROOT / "bin" / "mst")], (ROOT / "tests" / "mst_sample.txt").read_text())

    (RESULTS / "selection_sample.out").write_text(sample_selection, encoding="utf-8")
    (RESULTS / "partition_sample.out").write_text(sample_partition, encoding="utf-8")
    (RESULTS / "mst_sample.out").write_text(sample_mst, encoding="utf-8")

    document = Document()
    set_document_fonts(document)
    title = document.add_heading("《数据结构与算法设计2》期中大作业报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("题目：分治算法、动态规划、贪心算法综合实现")
    document.add_paragraph("学号：学号待填写    姓名：姓名待填写")
    document.add_paragraph("编程语言及环境：C++11，Apple clang 21.0.0，CMake 4.3.1；报告由 python-docx 生成。")

    document.add_heading("一、分治算法：第 k 小元素", level=1)
    document.add_paragraph(
        "问题要求在无序线性序集中找出第 k 小元素。程序实现了三种方法：容量为 k 的最大堆、随机划分线性选择、以及中位数的中位数线性选择。"
    )
    document.add_paragraph(
        "最大堆方法先用前 k 个元素建堆，之后扫描其余元素，只在新元素小于堆顶时替换，时间复杂度为 O(k + (n-k)logk)，空间复杂度为 O(k)。"
    )
    document.add_paragraph(
        "随机划分选择使用三路划分处理重复值，平均期望时间复杂度为 O(n)，最坏情况为 O(n^2)。中位数的中位数方法将元素按 5 个一组取中位数，再递归选择枢轴，保证最坏时间复杂度为 O(n)。"
    )
    document.add_heading("样例运行结果", level=2)
    add_code_block(document, sample_selection)
    document.add_heading("运行时间随 k 变化", level=2)
    add_report_table(
        document,
        ["n", "k", "堆选择/us", "随机选择/us", "中位数选择/us"],
        [
            [
                str(row["n"]),
                str(row["k"]),
                str(row["heap_us"]),
                str(row["randomized_us"]),
                str(row["median_of_medians_us"]),
            ]
            for row in by_k_rows
        ],
    )
    document.add_picture(str(RESULTS / "selection_by_k.png"), width=Inches(5.7))
    document.add_heading("运行时间随 n 变化", level=2)
    add_report_table(
        document,
        ["n", "k", "堆选择/us", "随机选择/us", "中位数选择/us"],
        [
            [
                str(row["n"]),
                str(row["k"]),
                str(row["heap_us"]),
                str(row["randomized_us"]),
                str(row["median_of_medians_us"]),
            ]
            for row in by_n_rows
        ],
    )
    document.add_picture(str(RESULTS / "selection_by_n.png"), width=Inches(5.7))
    document.add_paragraph(
        "实验中随机划分方法通常常数更小；中位数的中位数方法具有更强的最坏情况保证，但分组、排序和递归取枢轴带来了额外常数开销。"
    )

    document.add_heading("二、动态规划：卡牌价值二分", level=1)
    document.add_paragraph(
        "该问题要求将 n 张正整数价值卡牌分成两堆，使两堆价值和差值最小。暴力法枚举全部子集，时间复杂度为 O(n2^n)，可作为小规模正确性基准。"
    )
    document.add_paragraph(
        "动态规划设 dp[i][s] 表示使用前 i 张卡牌是否能组成价值 s。目标只需搜索不超过总价值一半的最大可达 s，再回溯得到第一堆卡牌编号。时间复杂度为 O(nS)，空间复杂度为 O(nS)，其中 S 为总价值的一半。"
    )
    document.add_heading("样例运行结果", level=2)
    add_code_block(document, sample_partition)
    document.add_heading("暴力法与动态规划实验", level=2)
    add_report_table(
        document,
        ["n", "总价值", "暴力/us", "动态规划/us", "最小差值"],
        [
            [
                str(row["n"]),
                str(row["total_value"]),
                str(row["brute_force_us"]),
                str(row["dynamic_programming_us"]),
                str(row["best_difference"]),
            ]
            for row in partition_rows
        ],
    )
    document.add_picture(str(RESULTS / "partition_experiments.png"), width=Inches(5.7))

    document.add_heading("三、贪心算法：供水管网布局", level=1)
    document.add_paragraph(
        "供水管网需要连通所有楼栋且不能出现环路，并要求总管道长度最小，本质是无向连通带权图的最小生成树问题。程序采用 Kruskal 算法：按边权从小到大扫描，用并查集判断是否成环，直到选择 n-1 条边。"
    )
    document.add_paragraph(
        "设楼栋数为 n、可铺设管道数为 m，排序复杂度为 O(mlogm)，并查集近似为 O(m alpha(n))，总复杂度为 O(mlogm)。"
    )
    document.add_heading("样例运行结果", level=2)
    add_code_block(document, sample_mst)

    document.add_heading("四、其他说明与体会", level=1)
    document.add_paragraph(
        "实现过程中需要特别处理重复元素、动态规划回溯、以及非连通图等边界情况。第 k 小元素选择中使用三路划分可以避免重复值导致递归区间缩小过慢；卡牌分堆中先求接近总价值一半的子集可直接保证两堆差值最小；最小生成树中若最终边数不足 n-1，则说明管网无法覆盖所有楼栋。"
    )
    document.add_paragraph(
        "通过实验可以看到，理论复杂度与实测趋势基本一致：暴力枚举随 n 增长迅速变慢，动态规划适合总价值规模可控的输入；随机选择在通常输入上效率较高，中位数的中位数方法则提供确定性的线性时间上界。"
    )

    document.save(ROOT / "report.docx")


def create_package() -> None:
    PACKAGE.mkdir(exist_ok=True)
    zip_path = PACKAGE / "学号姓名期中大作业.zip"
    include_paths = [
        ROOT / "report.docx",
        ROOT / "题目讲解.md",
        ROOT / "CMakeLists.txt",
        ROOT / "src",
        ROOT / "tests",
        ROOT / "scripts",
        ROOT / "results",
    ]
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path in include_paths:
            if path.is_file():
                archive.write(path, path.relative_to(ROOT))
            else:
                for child in sorted(path.rglob("*")):
                    if child.is_file():
                        archive.write(child, child.relative_to(ROOT))


def main() -> None:
    RESULTS.mkdir(exist_ok=True)
    PACKAGE.mkdir(exist_ok=True)
    build_project()
    by_k_rows, by_n_rows, partition_rows = generate_experiments()
    generate_report(by_k_rows, by_n_rows, partition_rows)
    create_package()


if __name__ == "__main__":
    main()
